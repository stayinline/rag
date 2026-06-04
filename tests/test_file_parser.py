"""Tests for file parser service."""
import os
import tempfile

from app.services.file_parser import (
    parse_pdf,
    parse_docx,
    parse_txt,
    parse_markdown,
    parse_file,
    compute_file_hash,
)


def test_parse_txt():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
        f.write("Hello, this is test content.\nSecond line.")
        f.flush()
        text = parse_txt(f.name)
    assert "Hello, this is test content" in text
    assert "Second line" in text
    os.unlink(f.name)


def test_parse_markdown():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
        f.write("# Title\n\nContent here.")
        f.flush()
        text = parse_markdown(f.name)
    assert "# Title" in text
    assert "Content here" in text
    os.unlink(f.name)


def test_parse_file_txt():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
        f.write("Test content")
        f.flush()
        text = parse_file(f.name)
    assert text == "Test content"
    os.unlink(f.name)


def test_parse_file_md():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
        f.write("# Header\nBody text")
        f.flush()
        text = parse_file(f.name)
    assert "# Header" in text
    os.unlink(f.name)


def test_parse_file_unsupported():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".xyz", delete=False) as f:
        f.write("content")
        f.flush()
        try:
            parse_file(f.name)
            assert False, "Should have raised ValueError"
        except ValueError as e:
            assert "Unsupported file type" in str(e)
    os.unlink(f.name)


def test_compute_file_hash():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
        f.write("test content")
        f.flush()
        hash1 = compute_file_hash(f.name)
        hash2 = compute_file_hash(f.name)
    assert hash1 == hash2
    assert len(hash1) == 64  # SHA-256 hex length
    os.unlink(f.name)


def test_compute_file_hash_different_content():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f1:
        f1.write("content A")
        f1.flush()
        hash_a = compute_file_hash(f1.name)
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f2:
        f2.write("content B")
        f2.flush()
        hash_b = compute_file_hash(f2.name)
    assert hash_a != hash_b
    os.unlink(f1.name)
    os.unlink(f2.name)


def test_parse_docx():
    from docx import Document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        doc.add_paragraph("Hello from DOCX")
        doc.add_paragraph("Second paragraph")
        doc.save(f.name)
        text = parse_docx(f.name)
    assert "Hello from DOCX" in text
    assert "Second paragraph" in text
    os.unlink(f.name)


def test_parse_file_docx():
    from docx import Document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        doc.add_paragraph("Test DOCX content")
        doc.save(f.name)
        text = parse_file(f.name)
    assert "Test DOCX content" in text
    os.unlink(f.name)


def test_parse_pdf():
    """Test PDF parsing with PyMuPDF."""
    import fitz
    fd, pdf_path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    try:
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Hello from PDF")
        doc.save(pdf_path)
        doc.close()

        text = parse_pdf(pdf_path)
        assert "Hello from PDF" in text
    finally:
        os.unlink(pdf_path)


def test_parse_file_pdf():
    """Test parse_file with PDF."""
    import fitz
    fd, pdf_path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    try:
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "PDF test content")
        doc.save(pdf_path)
        doc.close()

        text = parse_file(pdf_path)
        assert "PDF test content" in text
    finally:
        os.unlink(pdf_path)
